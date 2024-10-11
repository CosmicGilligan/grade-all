from operator import itemgetter
import os
from operator import itemgetter
import pandas as pd
import numpy as np
import tiktoken
from ast import literal_eval
from scipy.spatial.distance import cosine
from bs4 import BeautifulSoup
import openai
from docx import Document
from pypdf import PdfReader
import requests
import json
import re

# Your Canvas API URL and token
API_URL = 'https://sdccd.instructure.com/api/v1'
API_TOKEN = '1069~JHu89fDn3RfmMLCcP87FfXENWKMfkk3B9Y4hK7PaNQf3RGtrvDMPLCECeYZZQ846'
#course_id = '71354'  # Replace with the actual course ID

# Headers for the request
headers = {
    'Authorization': f'Bearer {API_TOKEN}'
}

client=openai

domain = "../db/Transcripts/"
subdomain109 = "../db/Transcripts/109/"
subdomain110 = "../db/Transcripts/110/"
full_url = "../db/Transcripts/"
suburl109 = "../db/Transcripts/109"
suburl110 = "../db/Transcripts/110"
max_tokens = 500
lecStrList = []
shortened = []
df = []
df_embeddings = []
df_similarities = []
prmtTitleList=[]
qFileList109=[]
qFileList110=[]
US1List=[]
US2List=[]
gradeMode=True
examMode=False
directory = './submissions/'


def extract_content(html_content):
    # Parse the HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Extract name enclosed in <h1>
    name = soup.find('h1').get_text(strip=True)
    
    # Initialize a variable to hold the extracted text
    extracted_text = ''
    
    # Find the <h1> tag, then start searching for subsequent text containers
    for sibling in soup.find('h1').next_siblings:
        if sibling.name in ['p', 'ol', 'div']:  # Add any other tags you expect to contain the text
            extracted_text += sibling.get_text(" ", strip=True) + ' '
        # Optional: Break out of the loop if a certain tag is encountered indicating the end of the relevant text
    
    # Clean up the extracted text
    extracted_text = extracted_text.strip()
    
    return name, extracted_text

def read_docx(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    newText = '\n'.join(full_text)
    text=newText.replace('\n', ' ')   
    return text
#    doc = doc = aw.Document(file_path)
#    return doc

def read_pdf(file_path):
    # creating a pdf reader object 
    reader = PdfReader(file_path) 
    
    text = ""
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        newText = page.extract_text()
        text+=newText.replace('\n', ' ')   
    return text

def read_file_content(file_path):
    if file_path.endswith('.docx'):
        return read_docx(file_path)
    elif file_path.endswith('.pdf'):
        return read_pdf(file_path)
    else:
        raise ValueError("Unsupported file format. Please use .docx or .pdf files.")
    
def getCanvasAPI(url):  
    params = {
        'per_page': 10000  # Setting the item count to 10 per page
    }

    # Make the GET request to retrieve courses
    response = requests.get(url, headers=headers)
    results = []
    while url and len(results) < 20:
            response = requests.get(url, headers=headers, params=params)
            if response.status_code != 200:
                raise Exception(f"API request failed with status code {response.status_code}")
            results.extend(response.json())
            # Check if there is a next page
            url = response.links.get('next', {}).get('url')
    #return results[:max_items]

    if response.status_code == 200:
        return results
    else:
        print(f"API request failed with status code {response.status_code}")
        return None

def getUserList(crn):
    # URL for listing all courses with additional parameters
    url = f'{API_URL}/courses/'+crn+'/users'
    students=getCanvasAPI(url)
    
    # Create a dictionary to map IDs to names
    id_to_name = {student['id']: student['name'] for student in students}

    # Directory containing the files
    directory = './submissions/'

    # List to store the name, filename pairs
    name_file_pairs = []

    # List all files in the directory
    files = os.listdir(directory)

    # Regular expression to extract ID from filename
    pattern = re.compile(r'_(\d+)_')

    for filename in files:
        match = pattern.search(filename)
        if match:
            student_id = int(match.group(1))
            if student_id in id_to_name:
                name_file_pairs.append((id_to_name[student_id], filename))

    # Print the result
    for name, filename in name_file_pairs:
        print(f'Name: {name}, Filename: {filename}')
    return name_file_pairs

def makeEntryList(crn):
    entryList = []
    userList = []
    gotUsers=False

    for filename in os.listdir(directory):
        if filename.endswith(('.html', '.docx', '.pdf')):
            filepath = os.path.join(directory, filename)
            
            if filename.endswith('.html'):
                with open(filepath, 'r', encoding='utf-8') as file:
                    thisEntry = []
                    html_content = file.read()
                    name, text = extract_content(html_content)
                    i = name.find(":")
                    strEnd = len(name)
                    shortName = name[i + 2:strEnd]
                    thisEntry.append(shortName)
                    thisEntry.append(text)
                    thisEntry.append("\n")
                    entryList.append(thisEntry)
            
            elif filename.endswith(('.pdf', '.docx')):
                thisEntry = []
                if gotUsers == False:
                    userList=getUserList(crn)
                    gotUsers=True
                # Create the entry list

                for name, filename in userList:
                    filepath = os.path.join(directory, filename)
                    data = read_file_content(filepath)
                    entryList.append([name, data, None])

    # Print the entry list
    for entry in entryList:
        print(f'Name: {entry[0]}, Data: {entry[1]}, Empty: {entry[2]}')

    return entryList

# search embedded docs based on cosine similarity

def calculate_similarity(embedding1, embedding2):
    # Using 1 - cosine distance to get cosine similarity
    result = 1 - cosine(embedding1, embedding2)
    return result

def get_embedding(text, model="text-embedding-ada-002"):
   return client.embeddings.create(input = [text], model='text-embedding-ada-002').data[0].embedding
              
def search_docs(df, user_query, top_n=3, to_print=True):
    embedding = get_embedding(
        user_query,
        model="text-embedding-ada-002"
    )
    
    df_similarities=df.copy()
    df_similarities["similarities"] = df.embeddings.apply(lambda x: calculate_similarity(embedding, x))

    res = (
        df_similarities.sort_values("similarities", ascending=False)
        .head(top_n)
    )
    if to_print:
        print(res)
    return res

def split_into_many(tokenizer, text, max_tokens = max_tokens):

    # Split the text into sentences
    sentences = text.split('. ')

    # Get the number of tokens for each sentence
    n_tokens = [len(tokenizer.encode(" " + sentence)) for sentence in sentences]
    
    chunks = []
    tokens_so_far = 0
    chunk = []

    # Loop through the sentences and tokens joined together in a tuple
    for sentence, token in zip(sentences, n_tokens):

        # If the number of tokens so far plus the number of tokens in the current sentence is greater 
        # than the max number of tokens, then add the chunk to the list of chunks and reset
        # the chunk and tokens so far
        if tokens_so_far + token > max_tokens:
            chunks.append(". ".join(chunk) + ".")
            chunk = []
            tokens_so_far = 0

        # If the number of tokens in the current sentence is greater than the max number of 
        # tokens, go to the next sentence
        if token > max_tokens:
            continue

        # Otherwise, add the sentence to the chunk and add the number of tokens to the total
        chunk.append(sentence)
        tokens_so_far += token + 1
        
    # Add the last chunk to the list of chunks
    if chunk:
        chunks.append(". ".join(chunk) + ".")

    return chunks

def remove_newlines(serie):
    serie = serie.str.replace('\n', ' ')
    serie = serie.str.replace('\\n', ' ')
    serie = serie.str.replace('  ', ' ')
    serie = serie.str.replace('  ', ' ')
    return serie

# Create a dataframe from the list of texts
def createDataframe():
    global shortened
    global df
    global df_embeddings
    global df_similarities

    df = pd.DataFrame(lecStrList, columns = ['fname', 'text', 'id'])
    df.to_csv('embeddings.csv')
    df.head()

    # Load the cl100k_base tokenizer which is designed to work with the ada-002 model
    tokenizer = tiktoken.get_encoding("cl100k_base")

    df = pd.read_csv('embeddings.csv', index_col=0)
    df.columns = ['title', 'text', 'id']

    # Tokenize the text and save the number of tokens to a new column
    df['n_tokens'] = df.text.apply(lambda x: len(tokenizer.encode(x)))

    # Visualize the distribution of the number of tokens per row using a histogram
    df.n_tokens.hist()

    # Loop through the dataframe
    for row in df.iterrows():

        # If the text is None, go to the next row
        if row[1].text is None:
            continue

        # If the number of tokens is greater than the max number of tokens, split the text into chunks
        if row[1].n_tokens > max_tokens:
            shortened += split_into_many(tokenizer,row[1].text)
        
        # Otherwise, add the text to the list of shortened texts
        else:
            shortened.append(row[1].text)

    df = pd.DataFrame(shortened, columns = ['text'])
    df['n_tokens'] = df.text.apply(lambda x: len(tokenizer.encode(x)))
    df.n_tokens.hist()

    # Note that you may run into rate limit issues depending on how many files you try to embed
    # Please check out our rate limit guide to learn more on how to handle this: https://platform.openai.com/docs/guides/rate-limits
#    df['embeddings'] = df.text.apply(lambda x: client.embeddings.create(input=x, engine='text-embedding-ada-002')['data'][0]['embedding'])
    df['embeddings'] = df.text.apply(lambda x: client.embeddings.create(input = x, model='text-embedding-ada-002').data[0].embedding)
                                               
    df.to_csv('tokens.csv')
    df.head()

    df=pd.read_csv('tokens.csv', index_col=0)
    df['embeddings'] = df['embeddings'].apply(literal_eval).apply(np.array)

    df.head()

def sort_by_number(string_list):

  # Function to extract the numeric part (xxx) from a string
  def get_number(text):
    return int(text.split(':')[-1])

  # Sort the list using the extracted numeric part as the key
  return sorted(string_list, key=get_number)

def crawl():
#    ospath=os.fspath(full_url)
#    osdir=os.listdir(full_url)
    dirTree=[]
#    max_tokens = 500
#    max_len = 1800
#    shortened = []
#    local_domain = domain
#    queue=[]
#    seen=set([])
    global qFileList109
    global qFileList110
    global lecStrList
    global US1List
    global US2List

    for (root,dirs,files) in os.walk(domain, topdown=True): 
        rootcrs=root.find("109")
        if(rootcrs>=0):
            rootdir=109
        rootcrs=root.find("110")
        if(rootcrs>=0):
            rootdir=110
        dirTree.append(dirs)      
        while files:
            foundQ=False
            fileList=files.pop(0)
            x=fileList.find("q.txt")
            if(x>=0):
                newFile="./"+root+"/"+fileList
                f=open(newFile)
                qinFileList=f.readlines()
                f.close()
                if(rootdir==109):
                    qFileList109.append(qinFileList)
                elif (rootdir==110):
                    qFileList110.append(qinFileList)
                foundQ=True
            x=fileList.find(".txt")
            if(( x>=0) and (foundQ==False)):
                newFile="./"+root+"/"+fileList
                titleList=[newFile]
                f=open(newFile, "r")
                inputList=(f.readlines())
                f.close
                lStr=inputList[2]
                l1Str=inputList[1].replace("\n","")
                lnStr=l1Str+lStr[:3]
                idnumstr=lStr[:3]
                x=inputList[1].find("US1")
                if(x>=0):                    
                    US1List.append(lnStr)
                else:
                    US2List.append(lnStr)

                removeStr=(inputList[1].replace("\n","")+inputList[2].replace("\n",""))
                idStr=inputList[2].replace("\n","")
                idStr1=idStr.replace("-","")
                idStr2=idStr1.replace(" ","")
                id=int(idStr2)
                titleList.append(id)
                titleStr=removeStr+inputList[3]
                lecStr=inputList[4].replace("\n","")
                tlecStrList=[]
                tlecStrList.append(titleStr)
                tlecStrList.append(lecStr)
                tlecStrList.append(idnumstr)
                lecStrList.append(tlecStrList)

    sorted_list=sorted(lecStrList,key=itemgetter(2))
    lecStrList=sorted_list
    sorted_list=sorted(qFileList109,key=itemgetter(0))
    qFileList109=sorted_list
    sorted_list=sorted(qFileList110,key=itemgetter(0))
    qFileList110=sorted_list
    sorted_list=sort_by_number(US1List)
    US1List=sorted_list
    sorted_list=sort_by_number(US2List)
    US2List=sorted_list
    createDataframe()

model="US1"
sStr=""

def getReviewPromptStr():
    print("Begin review completions...\n")
    filename=full_url+"rvwquery.prmt"
    f=open(filename)
    promptList=f.readlines()
    f.close()
    return promptList[0]

def getPromptStr(course,base):
    print("Begin grading completions...\n")

    #get the prompts

    filename=full_url+"query.prmt"
    f=open(filename)
    promptList=f.readlines()
    f.close()
    prompt1=promptList[0].replace("\n"," ")
    prompt2=promptList[1].replace("\n"," ")

    #get the lecture

    if(course==109):
        courseStr="US1"
    else:
        courseStr="US2"
    itemStr = "{:03}".format(base)
    lecStr=""
    for sublist in lecStrList:
        x=sublist[0].find(courseStr)
        if (x>=0):
            y=sublist[0].find(itemStr)
            if(y>=0):
                lecStr=sublist[1].replace("\n"," ")

    # get the question(s)

    baseStr=str(base)
    if course == 109:
        qList=qFileList109
    else:
        qList=qFileList110

    qStr=""
    for sublist in qList:
        x=sublist[0].find(itemStr)
        if(x>=0):
            listlen=len(sublist)
            y=2
            while(y<listlen):
                nxtStr=sublist[y].replace("\n"," ")
                qStr=qStr+nxtStr
                y=y+1

#    for x in qFileList: #This code gets the questions
#        y=x[0].find(baseStr)
#        if y>=0:
#            filename=rootStr+x[1].replace("\n","")
#            f=open(filename,"r")
#            qStrList=f.readlines()
#            f.close()
#            qLen=len(qStrList)
#            qStr=""
#            q=2
#            while q<qLen:
#                qStr=qStr+qStrList[q]
#                q+=1


    # The main prompt, the lecture text, The instructions after the lecture, and the questions
    # This function returns everything in the prompt except the student submission 
    promptStr=prompt1+" "+lecStr+" "+prompt2+" "+qStr+"\n" 
    return(promptStr)

def main():
    import sys    
    crawl()
    makeEntryList()
    print("done")

if __name__ == "__main__":
    main()
